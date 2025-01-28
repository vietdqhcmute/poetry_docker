import re

from handler import add_page_number, add_table_of_content, doc_to_docx, format_path, insert_paragraph_before, update_document

from docx import Document
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt


def main():
    doc_path ='./doc/luat_dan_su.docx'
    doc_path = format_path(doc_path)
    _pdf = "FALSE"
    # _pdf = input("Nếu muốn tạo pdf thì gõ true: ")


    if doc_path.endswith(".doc"):
        doc_path = doc_to_docx(doc_path)

    doc = Document(doc_path)
    styles = doc.styles


    ### Xử lý heading
    heading_list = ["Heading 1", "Heading 2", "Heading 3"]
    for heading in heading_list:
        if heading not in styles:
            styles.add_style(heading, WD_STYLE_TYPE.PARAGRAPH, builtin = True)

    key_words = ["(^Điều.*[.].*)", "^Mục.*[.]$"]
    key_words_2_line = ["^Chương.*", "^BỘ LUẬT.*", "^LUẬT.*", "^NGHỊ ĐỊNH.*", "^Phần thứ.*"]

    count = 1
    for i in range(0, len(doc.paragraphs)):
        para = doc.paragraphs[i]
        for key_word in key_words:
            if re.match(key_word, para.text):
                para.style = styles["Heading 3"]
                para.style.font.size = Pt(14)
                print(count, para.text, para.style)
                count += 1
        for key_word in key_words_2_line:
            if re.match(key_word, para.text):
                para.style = styles["Heading 1"]
                para.style.font.size = Pt(16)
                doc.paragraphs[i+1].style = styles["Heading 2"]
                doc.paragraphs[i+1].style.font.size = Pt(16)
                print(count, para.text, para.style)
                count += 1

    #Thêm Mục lục
    for item in doc.iter_inner_content():
        para_clone = insert_paragraph_before(item, "")
        para_clone.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        # para_clone.style.font.size = Pt(12)
        add_table_of_content(para_clone)
        para_clone.add_run().add_break(WD_BREAK.PAGE)
        break

    #Thêm số trang
    add_page_number(doc)

    doc.save(doc_path)

    #Cập nhật mục lục và chuyển thành PDF
    update_document(doc_path, _pdf.upper())

main()
