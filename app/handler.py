import re
import os
import inspect

import docx.document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx import Document

def format_path(_path):
    if "& " in _path or "'" in _path or '"' in _path:
        _path = _path.replace("'","")
        _path = _path.replace('"',"")
        _path = _path.replace('& ',"")
    return _path

#Add_paragraph_before_any_item
def insert_paragraph_before(item, text, style=None):
    """
    Return a newly created paragraph, inserted directly before this
    item (Table, etc.).
    """

    p = CT_P.add_p_before(item._element)
    p2 = Paragraph(p, item._parent)
    p2.text = text
    p2.style = style
    return p2

#Add_Table of Content
def add_table_of_content(paragraph):
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')  # creates a new element
    # fldChar.set(qn('w:dirty'), 'true') #Make TOC auto update on openning
    fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'   # change 1-3 depending on heading levels you need

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = "Right-click to update field."
    fldChar2.append(fldChar3)

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)
    p_element = paragraph._p

def update_document(file_path, pdf = "False"):
    script_dir = os.path.dirname(os.path.abspath(inspect.getfile(inspect.currentframe())))
    file_path = os.path.join(script_dir, file_path)
    doc = Document(file_path)
    # doc.TablesOfContents(1).Update()
    if pdf == "TRUE":
        doc.SaveAs2(file_path.replace(".docx", ".pdf"), FileFormat = 17)
    doc.save('output.docx')

def create_run_xml(run):
    fldStart = OxmlElement('w:fldChar')
    fldStart.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"

    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'separate')

    fldChar2 = OxmlElement('w:t')
    fldChar2.text = "2"

    fldEnd = OxmlElement('w:fldChar')
    fldEnd.set(qn('w:fldCharType'), 'end')

    run._r.append(fldStart)
    run._r.append(instrText)
    run._r.append(fldChar1)
    run._r.append(fldChar2)
    run._r.append(fldEnd)

def add_page_number(doc):
    for section in doc.sections:
        create_run_xml(section.footer.add_paragraph().add_run())
        sectPr = section._sectPr

        pgNumType = OxmlElement('w:pgNumType')
        pgNumType.set(qn('w:start'), "0")
        sectPr.append(pgNumType)

def doc_to_docx(doc_path):
    word = Document(doc_path)
    word.visible = 0
    wb = word.Documents.Open(doc_path)
    out_file = doc_path.replace(".doc", ".docx")
    wb.SaveAs2(out_file, FileFormat=16) # file format for docx
    wb.Close()
    word.Quit()
    return out_file
